import os
from datetime import datetime
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

load_dotenv()
client = OpenAI()

GLOBAL_VECTOR_STORE = "vs_69b30bc93c1081919e375324427ccc05"
PROJECT_VECTOR_STORE = "vs_69b4035175c48191ae87cbfcf0663248"

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def save_docx(text: str, title: str, filename: str) -> str:
    path = os.path.join(OUTPUT_DIR, filename)
    doc = Document()
    doc.add_heading(title, level=1)

    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)

    doc.save(path)
    return path


def save_pdf(text: str, title: str, filename: str) -> str:
    path = os.path.join(OUTPUT_DIR, filename)

    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    x_margin = 40
    y = height - 50
    max_width = width - 2 * x_margin

    c.setFont("Helvetica-Bold", 14)
    c.drawString(x_margin, y, title)
    y -= 30

    c.setFont("Helvetica", 10)

    for paragraph in text.split("\n"):
        lines = simpleSplit(paragraph, "Helvetica", 10, max_width)

        if not lines:
            y -= 12

        for line in lines:
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 10)

            c.drawString(x_margin, y, line)
            y -= 12

    c.save()
    return path


def should_generate_files(user_input: str) -> bool:
    triggers = [
        "gere", "gerar", "crie", "criar", "template", "templates",
        "docx", "pdf", "relatório", "relatorio", "checklist",
        "formulário", "formulario", "word", "excel", "planilha"
    ]
    text = user_input.lower()
    return any(t in text for t in triggers)


def build_title_from_question(user_input: str) -> str:
    text = user_input.lower()

    if "checklist" in text:
        return "Checklist gerado pela AiA"
    if "template" in text or "templates" in text:
        return "Templates gerados pela AiA"
    if "relatório" in text or "relatorio" in text:
        return "Relatório gerado pela AiA"
    return "Documento gerado pela AiA"


def ask_ai(question: str) -> str:
    prompt = f"""
Você é a AiA — AstraCarbon Inteligência Artificial.

Funções:
- analisar conformidade metodológica de projetos de biochar;
- comparar metodologia com documentos do projeto;
- estruturar lacunas, checklists, templates e relatórios;
- responder sempre em português do Brasil;
- ser objetivo, técnico e organizado.

Regras:
- use a base metodológica para identificar o que a metodologia exige;
- use a base do projeto para verificar o que o projeto apresenta;
- quando houver lacunas, aponte diretamente o documento faltante;
- quando o usuário pedir checklist, relatório ou template, entregue em formato utilizável.

Pergunta do usuário:
{question}
"""

    response = client.responses.create(
        model="gpt-5-mini",
        input=prompt,
        tools=[
            {
                "type": "file_search",
                "vector_store_ids": [
                    GLOBAL_VECTOR_STORE,
                    PROJECT_VECTOR_STORE
                ]
            }
        ]
    )

    return response.output_text


def show_menu():
    print("\n" + "=" * 60)
    print("AiA — AstraCarbon Inteligência Artificial")
    print("=" * 60)
    print("1. Fazer pergunta")
    print("2. Gerar checklist de lacunas metodológicas")
    print("3. Gerar relatório de conformidade")
    print("4. Sair")
    print("=" * 60)


def main():
    while True:
        show_menu()
        choice = input("Escolha uma opção: ").strip()

        if choice == "1":
            question = input("\nPergunta: ").strip()

        elif choice == "2":
            question = (
                "Gere um checklist dos documentos faltantes e das lacunas "
                "metodológicas do projeto, comparando PDD e ACV com a metodologia."
            )

        elif choice == "3":
            question = (
                "Gere um relatório de conformidade metodológica do projeto, "
                "comparando os requisitos da metodologia com as evidências do PDD e do ACV."
            )

        elif choice == "4":
            print("Encerrando a AiA.")
            break

        else:
            print("Opção inválida.")
            continue

        print("\nProcessando...\n")
        text = ask_ai(question)
        print(text)

        if should_generate_files(question):
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            title = build_title_from_question(question)

            docx_path = save_docx(text, title, f"aia_{timestamp}.docx")
            pdf_path = save_pdf(text, title, f"aia_{timestamp}.pdf")

            print("\nArquivos gerados:")
            print(docx_path)
            print(pdf_path)


if __name__ == "__main__":
    main()